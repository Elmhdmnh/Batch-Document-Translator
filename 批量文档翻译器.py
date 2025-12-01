#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import threading
import queue
import time
import traceback
from tkinter import Tk, Button, Label, Entry, StringVar, filedialog, ttk, Text, END, W
import docx2txt
from docx import Document
import requests

MAX_CHUNK_CHARS = 15000
RETRY_COUNT = 3
RETRY_DELAY = 2


class TranslatorGUI:
    def __init__(self, root):
        self.root = root
        root.title("批量文档翻译器")
        root.geometry("860x620")

        self.api_key_var = StringVar()
        self.base_url_var = StringVar(value="https://api.openai.com")
        self.model_var = StringVar(value="gpt-4o-mini")

        # 新增：用户输入目标语言 + 风格
        self.target_lang_var = StringVar(value="中文")
        self.style_var = StringVar(value="信达雅")

        self.output_dir_var = StringVar(value=os.getcwd())
        self.file_list = []
        self.stop_flag = threading.Event()

        self._build_ui()

    def _build_ui(self):
        Label(self.root, text="API Key:").grid(row=0, column=0, sticky=W, padx=8, pady=6)
        Entry(self.root, textvariable=self.api_key_var, width=60, show='*').grid(row=0, column=1, columnspan=3, sticky=W)

        Label(self.root, text="Base URL:").grid(row=1, column=0, sticky=W, padx=8, pady=6)
        Entry(self.root, textvariable=self.base_url_var, width=60).grid(row=1, column=1, columnspan=3, sticky=W)

        Label(self.root, text="Model:").grid(row=2, column=0, sticky=W, padx=8, pady=6)
        Entry(self.root, textvariable=self.model_var, width=30).grid(row=2, column=1, sticky=W)

        # ★ 新增：目标语言
        Label(self.root, text="目标语言:").grid(row=3, column=0, sticky=W, padx=8, pady=6)
        Entry(self.root, textvariable=self.target_lang_var, width=20).grid(row=3, column=1, sticky=W)

        # ★ 新增：翻译风格
        Label(self.root, text="翻译风格:").grid(row=3, column=2, sticky=W, padx=8, pady=6)
        Entry(self.root, textvariable=self.style_var, width=20).grid(row=3, column=3, sticky=W)

        Label(self.root, text="输出目录:").grid(row=4, column=0, sticky=W, padx=8, pady=6)
        Entry(self.root, textvariable=self.output_dir_var, width=60).grid(row=4, column=1, columnspan=2, sticky=W)
        Button(self.root, text="选择目录", command=self.choose_output_dir).grid(row=4, column=3, sticky=W)

        Button(self.root, text="选择文件（多选）", command=self.select_files, width=20).grid(row=5, column=0, padx=8, pady=10)
        Button(self.root, text="开始翻译", command=self.start_translation, width=20).grid(row=5, column=1, padx=8)
        Button(self.root, text="停止", command=self.stop_translation, width=12).grid(row=5, column=2)
        Button(self.root, text="清空日志", command=self.clear_log, width=12).grid(row=5, column=3)

        self.progress = ttk.Progressbar(self.root, orient='horizontal', length=700, mode='determinate')
        self.progress.grid(row=6, column=0, columnspan=4, padx=8, pady=6)

        Label(self.root, text="已选择文件:").grid(row=7, column=0, sticky=W, padx=8)
        self.files_text = Text(self.root, height=6, width=100)
        self.files_text.grid(row=8, column=0, columnspan=4, padx=8)

        Label(self.root, text="日志:").grid(row=9, column=0, sticky=W, padx=8)
        self.log_text = Text(self.root, height=12, width=100)
        self.log_text.grid(row=10, column=0, columnspan=4, padx=8, pady=8)

    def choose_output_dir(self):
        d = filedialog.askdirectory()
        if d:
            self.output_dir_var.set(d)

    def select_files(self):
        paths = filedialog.askopenfilenames(title="选择 文档", filetypes=[("文档", "*.docx *.doc *.txt")])
        if not paths:
            return
        self.file_list = list(paths)
        self.root.after(0, lambda: self._update_file_list())
        self.log(f"选择了 {len(self.file_list)} 个文件")

    def _update_file_list(self):
        self.files_text.delete(1.0, END)
        for p in self.file_list:
            self.files_text.insert(END, p + "\n")

    def log(self, *parts):
        t = time.strftime('%H:%M:%S')
        msg = f"[{t}] " + " ".join(map(str, parts)) + "\n"
        self.root.after(0, lambda: self._append_log(msg))

    def _append_log(self, msg):
        self.log_text.insert(END, msg)
        self.log_text.see(END)

    def clear_log(self):
        self.log_text.delete(1.0, END)

    def stop_translation(self):
        self.stop_flag.set()
        self.log("用户请求停止。")

    def start_translation(self):
        if not self.file_list:
            self.log("错误：未选择文件")
            return
        self.stop_flag.clear()
        self.progress["maximum"] = len(self.file_list)
        self.progress["value"] = 0
        threading.Thread(target=self._worker_thread, daemon=True).start()

    def _worker_thread(self):
        total = len(self.file_list)
        for idx, path in enumerate(self.file_list, start=1):
            if self.stop_flag.is_set():
                break
            try:
                self.log(f"开始翻译：{os.path.basename(path)} ({idx}/{total})")
                self._translate_file(path)
                self.log(f"完成：{os.path.basename(path)}")
            except Exception as e:
                self.log(f"失败：{e}")
                traceback.print_exc()
            finally:
                self.root.after(0, lambda v=idx: self.progress.configure(value=v))
        self.log("任务结束。")

    # -------------------- 文件读取 --------------------

    def _read_docx(self, filepath):
        try:
            doc = Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs)
        except:
            return docx2txt.process(filepath)

    def _read_doc(self, filepath):
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(filepath)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            return text
        except:
            raise ValueError("处理 .doc 需要 pywin32: pip install pywin32")

    def _read_txt(self, filepath):
        return open(filepath, "r", encoding="utf-8", errors="ignore").read()

    # -------------------- API 处理 --------------------

    def _prepare_endpoints(self, base_url):
        b = base_url.rstrip("/")
        if b.endswith("/v1"):
            b = b[:-3]
        return [
            f"{b}/v1/chat/completions",
            f"{b}/chat/completions",
            f"{b}/v1/responses",
            f"{b}/responses",
        ]

    def _chunks(self, text):
        if len(text) <= MAX_CHUNK_CHARS:
            return [text]
        parts, buf, blen = [], [], 0
        for p in text.split("\n"):
            l = len(p) + 1
            if blen + l > MAX_CHUNK_CHARS and buf:
                parts.append("\n".join(buf))
                buf, blen = [p], l
            else:
                buf.append(p)
                blen += l
        if buf:
            parts.append("\n".join(buf))
        return parts

    def _extract_translation(self, js):
        try:
            return js["choices"][0]["message"]["content"]
        except:
            pass
        try:
            return js["choices"][0]["text"]
        except:
            pass
        try:
            return js["output_text"]
        except:
            pass
        return None

    def _call_api(self, endpoints, payload, headers):
        last = None
        for ep in endpoints:
            for _ in range(RETRY_COUNT):
                try:
                    r = requests.post(ep, json=payload, headers=headers, timeout=120)
                    if r.status_code == 200:
                        return r.json()
                    last = r.text
                    time.sleep(RETRY_DELAY)
                except Exception as e:
                    last = e
                    time.sleep(RETRY_DELAY)
        raise RuntimeError(last)

    # -------------------- 翻译主逻辑 --------------------

    def _translate_file(self, filepath):
        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".docx":
            raw = self._read_docx(filepath)
        elif ext == ".doc":
            raw = self._read_doc(filepath)
        else:
            raw = self._read_txt(filepath)

        if not raw.strip():
            raise ValueError("文件内容为空")

        target = self.target_lang_var.get().strip()
        style = self.style_var.get().strip()

        # ★ 新 prompt：用户自由输入目标语言 + 风格
        system_prompt = (
            f"你是一名专业的翻译人员。\n"
            f"目标语言：{target}\n"
            f"翻译风格：{style}\n"
            "必须忠实原文，不增删信息。保持段落结构一致。"
        )

        endpoints = self._prepare_endpoints(self.base_url_var.get())
        headers = {"Content-Type": "application/json"}
        if self.api_key_var.get().strip():
            headers["Authorization"] = f"Bearer {self.api_key_var.get().strip()}"

        model = self.model_var.get().strip()
        chunks = self._chunks(raw)
        results = []

        for i, ck in enumerate(chunks, start=1):
            self.log(f"翻译分段 {i}/{len(chunks)}")
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": ck},
                ],
            }
            js = self._call_api(endpoints, payload, headers)
            txt = self._extract_translation(js)
            if not txt:
                raise RuntimeError("无法解析返回内容")
            results.append(txt)

        final = "\n".join(results)

        out_dir = self.output_dir_var.get()
        os.makedirs(out_dir, exist_ok=True)

        base = os.path.splitext(os.path.basename(filepath))[0]
        txt_path = os.path.join(out_dir, base + "_translated.txt")
        docx_path = os.path.join(out_dir, base + "_translated.docx")

        open(txt_path, "w", encoding="utf-8").write(final)

        doc = Document()
        for line in final.split("\n"):
            doc.add_paragraph(line)
        doc.save(docx_path)

        self.log(f"保存：{txt_path}")
        self.log(f"保存：{docx_path}")


if __name__ == "__main__":
    root = Tk()
    app = TranslatorGUI(root)
    root.mainloop()
